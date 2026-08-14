from pathlib import Path
import re

import spacy
from docx import Document


INPUT_FILE = Path("input/KSH_International_RHP.docx")
REDACTED_FILE = Path("output/KSH_International_RHP_REDACTED.docx")

MODEL_NAME = "en_core_web_sm"


def extract_document_text(path):
    document = Document(path)

    parts = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)


def show_context(text, value, radius=180):

    positions = []

    start = 0

    while True:

        index = text.lower().find(
            value.lower(),
            start
        )

        if index == -1:
            break

        positions.append(index)

        start = index + max(
            1,
            len(value)
        )

    for position in positions:

        left = max(
            0,
            position - radius
        )

        right = min(
            len(text),
            position + len(value) + radius
        )

        context = text[left:right]

        print()
        print("-" * 80)
        print(
            f"FOUND: {text[position:position + len(value)]}"
        )
        print("-" * 80)
        print(context.replace("\n", " "))
        print("-" * 80)


def main():

    print("=" * 80)
    print("PII CANDIDATE REVIEW")
    print("=" * 80)

    original_text = extract_document_text(
        INPUT_FILE
    )

    redacted_text = extract_document_text(
        REDACTED_FILE
    )

    nlp = spacy.load(
        MODEL_NAME
    )

    doc = nlp(redacted_text)

    persons = []
    organizations = []

    for entity in doc.ents:

        value = entity.text.strip()

        if not value:
            continue

        if entity.label_ == "PERSON":

            persons.append(value)

        elif entity.label_ == "ORG":

            organizations.append(value)

    persons = sorted(
        set(persons),
        key=str.lower
    )

    organizations = sorted(
        set(organizations),
        key=str.lower
    )

    print()
    print("=" * 80)
    print("PERSON CANDIDATES")
    print("=" * 80)

    for number, person in enumerate(
        persons,
        start=1
    ):

        print(
            f"{number:03d}. {person}"
        )

    print()
    print("=" * 80)
    print("ORGANIZATION CANDIDATES")
    print("=" * 80)

    for number, organization in enumerate(
        organizations,
        start=1
    ):

        print(
            f"{number:03d}. {organization}"
        )

    print()
    print("=" * 80)
    print("KNOWN IMPORTANT NAMES")
    print("=" * 80)

    important_names = [
        "Abhijit Diwan",
        "Ajay Menon",
        "Jayaram N. Shetty",
        "Karunakar Hegde",
        "Karunakar N. Bhandary",
        "Kumar Tiwari",
        "Kushal Hegde",
        "Kushal Subbayya Hegde",
        "Narayana B. Shetty",
        "Pushpa Kushal Hegde",
        "Rajesh Hegde",
        "Rajesh Kushal Hegde",
        "Rohit Hegde",
        "Rohit Kushal Hegde",
        "Rupal K. Sancheti",
        "Salil Ajay Bhargava",
        "Sangeeta Ramprasad Rai",
        "Shanti Gopalkrishnan",
        "Kishan Rastogi",
        "Sandesh Bhagwat",
        "Amod Joshi",
        "Ajay Shriram Patil",
        "Eric Bacha",
        "Sharmila Joshi",
        "Cherag Gyara",
        "Indu Jacob",
        "Vijay Hegde",
        "Kumar Tiwari",
        "Rakhi Girija Shetty",
        "Lokesh Shah",
        "Soumavo Sarkar",
        "Dinesh Hirachand Munot",
        "Lalit Muljibhai Sarvaiya",
        "Ganesh Prasad",
    ]

    for name in important_names:

        pattern = re.compile(
            r"(?<![A-Za-z])"
            + re.escape(name)
            + r"(?![A-Za-z])",
            re.IGNORECASE,
        )

        match = pattern.search(
            redacted_text
        )

        if match:

            print(
                f"\nSTILL PRESENT: {name}"
            )

            show_context(
                redacted_text,
                match.group()
            )

        else:

            print(
                f"REDACTED: {name}"
            )

    print()
    print("=" * 80)
    print("REVIEW COMPLETE")
    print("=" * 80)


if __name__ == "__main__":
    main()