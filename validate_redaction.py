from __future__ import annotations

import re
from pathlib import Path

import phonenumbers
from docx import Document


# ============================================================
# CONFIGURATION
# ============================================================

ORIGINAL_FILE = Path(
    "input/KSH_International_RHP.docx"
)

REDACTED_FILE = Path(
    "output/KSH_International_RHP_REDACTED.docx"
)


# ============================================================
# REGEX PATTERNS
# ============================================================

EMAIL_PATTERN = re.compile(
    r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"
)


PHONE_PATTERN = re.compile(
    r"(?<!\d)"
    r"(?:\+?\d[\d\s().-]{8,}\d)"
    r"(?!\d)"
)


IP_PATTERN = re.compile(
    r"\b(?:"
    r"(?:25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)\."
    r"){3}"
    r"(?:25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)\b"
)


CREDIT_CARD_PATTERN = re.compile(
    r"\b(?:"
    r"\d{4}[- ]?"
    r"){3}"
    r"\d{4}\b"
)


SSN_PATTERN = re.compile(
    r"\b\d{3}-\d{2}-\d{4}\b"
)


# ============================================================
# PHONE VALIDATION
# ============================================================

def valid_phone(candidate):

    digits = re.sub(
        r"\D",
        "",
        candidate,
    )

    # Phone numbers must have a reasonable length.
    if not 10 <= len(digits) <= 13:

        return False

    try:

        number = phonenumbers.parse(
            candidate,
            "IN",
        )

        return (
            phonenumbers.is_possible_number(
                number
            )
            and
            phonenumbers.is_valid_number(
                number
            )
        )

    except phonenumbers.NumberParseException:

        return False


# ============================================================
# DOCUMENT TEXT EXTRACTION
# ============================================================

def extract_text(document):

    parts = []

    # --------------------------------------------------------
    # PARAGRAPHS
    # --------------------------------------------------------

    for paragraph in document.paragraphs:

        text = paragraph.text

        if text.strip():

            parts.append(text)

    # --------------------------------------------------------
    # TABLES
    # --------------------------------------------------------

    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                text = cell.text

                if text.strip():

                    parts.append(text)

    return "\n".join(parts)


# ============================================================
# LOAD DOCUMENT
# ============================================================

def load_document(path):

    if not path.exists():

        raise FileNotFoundError(
            f"File not found: {path}"
        )

    return Document(path)


# ============================================================
# EXTRACT ORIGINAL PII
# ============================================================

def extract_sensitive_values(text):

    findings = {
        "email": set(),
        "phone": set(),
        "ip": set(),
        "credit_card": set(),
        "ssn": set(),
    }

    # --------------------------------------------------------
    # EMAIL
    # --------------------------------------------------------

    for match in EMAIL_PATTERN.finditer(text):

        findings["email"].add(
            match.group().strip().lower()
        )

    # --------------------------------------------------------
    # PHONE
    # --------------------------------------------------------

    for match in PHONE_PATTERN.finditer(text):

        candidate = match.group().strip()

        if valid_phone(candidate):

            findings["phone"].add(
                candidate
            )

    # --------------------------------------------------------
    # IP
    # --------------------------------------------------------

    for match in IP_PATTERN.finditer(text):

        findings["ip"].add(
            match.group().strip()
        )

    # --------------------------------------------------------
    # CREDIT CARD
    # --------------------------------------------------------

    for match in CREDIT_CARD_PATTERN.finditer(text):

        findings["credit_card"].add(
            match.group().strip()
        )

    # --------------------------------------------------------
    # SSN
    # --------------------------------------------------------

    for match in SSN_PATTERN.finditer(text):

        findings["ssn"].add(
            match.group().strip()
        )

    return findings


# ============================================================
# VALUE SURVIVAL CHECK
# ============================================================

def normalize_phone(value):

    return re.sub(
        r"\D",
        "",
        value,
    )


def phone_survives(
    original_phone,
    redacted_text,
):

    original_digits = normalize_phone(
        original_phone
    )

    # Check exact formatting first.
    if original_phone.lower() in redacted_text.lower():

        return True

    # Check normalized digits.
    #
    # This catches cases where formatting changed,
    # e.g.
    #
    # +91 20 1234 5678
    #
    # becomes
    #
    # 912012345678
    #
    # without falsely matching ordinary numbers.

    digit_runs = re.findall(
        r"\d{10,13}",
        redacted_text,
    )

    for digit_run in digit_runs:

        if digit_run == original_digits:

            return True

    return False


def value_survives(
    pii_type,
    value,
    redacted_text,
):

    if pii_type == "phone":

        return phone_survives(
            value,
            redacted_text,
        )

    return (
        value.lower()
        in redacted_text.lower()
    )


# ============================================================
# PLACEHOLDER COUNTS
# ============================================================

def count_placeholders(text):

    patterns = {

        "email": re.compile(
            r"person\d+@example\.com",
            re.IGNORECASE,
        ),

        "phone": re.compile(
            r"\+91 90000 \d{5}",
            re.IGNORECASE,
        ),

        "ip": re.compile(
            r"192\.0\.2\.\d+",
            re.IGNORECASE,
        ),

        "credit_card": re.compile(
            r"4111-1111-1111-\d{4}",
            re.IGNORECASE,
        ),

        "ssn": re.compile(
            r"000-00-\d{4}",
            re.IGNORECASE,
        ),

        "person": re.compile(
            r"Person \d+",
            re.IGNORECASE,
        ),

        "organization": re.compile(
            r"Example Organization \d+",
            re.IGNORECASE,
        ),

        "address": re.compile(
            r"Example Address \d+",
            re.IGNORECASE,
        ),
    }

    counts = {}

    for pii_type, pattern in patterns.items():

        counts[pii_type] = len(
            pattern.findall(text)
        )

    return counts


# ============================================================
# MAIN
# ============================================================

def main():

    print("=" * 70)
    print("ORIGINAL vs REDACTED DOCUMENT VALIDATION")
    print("=" * 70)

    print()

    print(
        f"Original : {ORIGINAL_FILE}"
    )

    print(
        f"Redacted : {REDACTED_FILE}"
    )

    print()

    # --------------------------------------------------------
    # LOAD DOCUMENTS
    # --------------------------------------------------------

    original_document = load_document(
        ORIGINAL_FILE
    )

    redacted_document = load_document(
        REDACTED_FILE
    )

    original_text = extract_text(
        original_document
    )

    redacted_text = extract_text(
        redacted_document
    )

    print(
        f"Original characters : "
        f"{len(original_text):,}"
    )

    print(
        f"Redacted characters : "
        f"{len(redacted_text):,}"
    )

    print()

    # --------------------------------------------------------
    # EXTRACT ORIGINAL PII
    # --------------------------------------------------------

    original_findings = extract_sensitive_values(
        original_text
    )

    # --------------------------------------------------------
    # SURVIVAL TEST
    # --------------------------------------------------------

    print("=" * 70)
    print("PII SURVIVAL TEST")
    print("=" * 70)

    print()

    total_original = 0
    total_survived = 0

    for pii_type, values in original_findings.items():

        total_original += len(values)

        survived = []

        for value in sorted(values):

            if value_survives(
                pii_type,
                value,
                redacted_text,
            ):

                survived.append(
                    value
                )

        total_survived += len(
            survived
        )

        print(
            f"{pii_type.upper():15}"
            f" original={len(values):4}"
            f" survived={len(survived):4}"
        )

        if survived:

            print()

            for value in survived:

                print(
                    f"    FAIL: {value}"
                )

        else:

            print(
                "    PASS"
            )

        print()

    # --------------------------------------------------------
    # PLACEHOLDER REPORT
    # --------------------------------------------------------

    print("=" * 70)
    print("GENERATED PLACEHOLDERS")
    print("=" * 70)

    print()

    placeholders = count_placeholders(
        redacted_text
    )

    for pii_type, count in placeholders.items():

        print(
            f"{pii_type:15}: {count}"
        )

    print()

    # --------------------------------------------------------
    # FINAL RESULT
    # --------------------------------------------------------

    print("=" * 70)
    print("FINAL VALIDATION RESULT")
    print("=" * 70)

    print()

    print(
        f"Original PII values checked : "
        f"{total_original}"
    )

    print(
        f"PII values surviving        : "
        f"{total_survived}"
    )

    print()

    if total_survived == 0:

        print(
            "STATUS: PASS"
        )

        print()

        print(
            "No detected high-confidence "
            "email, phone, IP, credit-card "
            "or SSN values from the original "
            "document survived."
        )

    else:

        print(
            "STATUS: FAIL"
        )

        print()

        print(
            "Some high-confidence sensitive "
            "values from the original document "
            "still exist in the redacted document."
        )

    print()

    print("=" * 70)
    print("VALIDATION COMPLETE")
    print("=" * 70)


# ============================================================
# ENTRY POINT
# ============================================================

if __name__ == "__main__":

    main()