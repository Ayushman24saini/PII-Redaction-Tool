from pathlib import Path

import spacy
from docx import Document


INPUT_FILE = Path("input/KSH_International_RHP.docx")
OUTPUT_FILE = Path("ner_inventory.txt")


def extract_text(document):
    chunks = []

    # Paragraphs
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            chunks.append(("paragraph", text))

    # Tables
    for table_index, table in enumerate(
        document.tables,
        start=1
    ):
        for row_index, row in enumerate(
            table.rows,
            start=1
        ):
            for col_index, cell in enumerate(
                row.cells,
                start=1
            ):
                text = cell.text.strip()

                if text:
                    location = (
                        f"table {table_index}, "
                        f"row {row_index}, "
                        f"column {col_index}"
                    )

                    chunks.append((location, text))

    # Headers and footers
    for section_index, section in enumerate(
        document.sections,
        start=1
    ):
        for paragraph in section.header.paragraphs:
            text = paragraph.text.strip()

            if text:
                chunks.append(
                    (
                        f"section {section_index} header",
                        text
                    )
                )

        for paragraph in section.footer.paragraphs:
            text = paragraph.text.strip()

            if text:
                chunks.append(
                    (
                        f"section {section_index} footer",
                        text
                    )
                )

    return chunks


def main():

    print("Loading spaCy model...")

    nlp = spacy.load("en_core_web_sm")

    document = Document(INPUT_FILE)

    chunks = extract_text(document)

    findings = []

    for location, text in chunks:

        doc = nlp(text)

        for entity in doc.ents:

            if entity.label_ in {
                "PERSON",
                "ORG",
                "GPE",
                "LOC",
            }:

                findings.append(
                    (
                        entity.label_,
                        location,
                        entity.text
                    )
                )

    with OUTPUT_FILE.open(
        "w",
        encoding="utf-8"
    ) as file:

        file.write("SPACY NER INVENTORY\n")
        file.write("=" * 70 + "\n\n")

        for label, location, value in findings:

            file.write(
                f"[{label}] "
                f"{location}: "
                f"{value}\n"
            )

    print("=" * 60)
    print("SPACY NER PROFILING COMPLETE")
    print("=" * 60)

    print(
        f"Text chunks scanned: {len(chunks)}"
    )

    print(
        f"Entities detected: {len(findings)}"
    )

    print(
        f"Detailed report: {OUTPUT_FILE}"
    )


if __name__ == "__main__":
    main()