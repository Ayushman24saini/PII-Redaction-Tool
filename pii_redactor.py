from __future__ import annotations

from operator import add
import re
from pathlib import Path

import phonenumbers
import spacy
from docx import Document

# ============================================================
# CONFIGURATION
# ============================================================

INPUT_FILE = Path(
    "input/KSH_International_RHP.docx"
)

OUTPUT_FILE = Path(
    "output/KSH_International_RHP_REDACTED.docx"
)

MODEL_NAME = "en_core_web_sm"


# ============================================================
# REGEX PATTERNS
# ============================================================

EMAIL_PATTERN = re.compile(
    r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"
)

IP_PATTERN = re.compile(
    r"\b(?:"
    r"(?:25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)\."
    r"){3}"
    r"(?:25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)\b"
)

CREDIT_CARD_PATTERN = re.compile(
    r"\b(?:"
    r"\d{4}[- ]?"
    r"){3}"
    r"\d{4}\b"
)

SSN_PATTERN = re.compile(
    r"\b\d{3}-\d{2}-\d{4}\b"
)

DOB_PATTERN = re.compile(
    r"""
    (?:
        date\s+of\s+birth |
        dob |
        d\.o\.b\. |
        born\s+on
    )
    \s*[:\-]?\s*
    (
        \d{1,2}[/-]\d{1,2}[/-]\d{2,4}
        |
        \d{1,2}\s+
        (?:January|February|March|April|May|June|
        July|August|September|October|November|December)
        \s+\d{4}
        |
        (?:January|February|March|April|May|June|
        July|August|September|October|November|December)
        \s+\d{1,2},?\s+\d{4}
    )
    """,
    re.IGNORECASE | re.VERBOSE,
)

PHONE_PATTERN = re.compile(
    r"(?<!\d)"
    r"(?:\+?\d[\d\s().-]{8,}\d)"
    r"(?!\d)"
)

PIN_PATTERN = re.compile(
    r"\b\d{6}\b"
)


# ============================================================
# ADDRESS DETECTION
# ============================================================

ADDRESS_LABEL_PATTERN = re.compile(
    r"""
    (?:
        registered\s+office |
        corporate\s+office |
        registered\s+address |
        corporate\s+address |
        mailing\s+address |
        residential\s+address |
        correspondence\s+address
    )
    \s*:
    """,
    re.IGNORECASE | re.VERBOSE,
)

ADDRESS_COMPONENT_PATTERN = re.compile(
    r"""
    (?:
        village |
        road |
        street |
        lane |
        tower |
        building |
        apartment |
        floor |
        taluka |
        district |
        industrial\s+area |
        nagar |
        farms |
        chakan |
        pune |
        mumbai |
        maharashtra |
        marg |
        complex |
        society |
        colony |
        nagar
    )
    """,
    re.IGNORECASE | re.VERBOSE,
)


# ============================================================
# PERSON DETECTION
# ============================================================

# Words that frequently appear in spaCy PERSON false positives.
GENERIC_PERSON_WORDS = {
    "offer",
    "offers",
    "promoters",
    "promoter",
    "directors",
    "director",
    "registrar",
    "company",
    "email",
    "website",
    "address",
    "amount",
    "bid",
    "bids",
    "bidder",
    "bidders",
    "bank",
    "banks",
    "shareholders",
    "shareholder",
    "investors",
    "investor",
    "members",
    "member",
    "board",
    "customers",
    "customer",
    "employees",
    "employee",
    "person",
    "reference",
    "rate",
    "floor",
    "price",
    "mutual",
    "funds",
    "fund",
    "air",
    "conditioning",
    "allotment",
    "allotted",
    "acknowledgement",
    "slip",
    "schedule",
    "key",
    "managerial",
    "personnel",
    "kmp",
    "upI",
    "upi",
    "circular",
    "circulars",
    "share",
    "shares",
    "transfer",
    "agents",
    "agent",
    "cap",
    "listing",
    "branch",
    "facility",
    "facilities",
    "defaulter",
    "wilful",
    "circuit",
    "kilometers",
    "gram",
    "jyoti",
    "mega",
    "volt",
    "amperes",
    "kisan",
    "urja",
    "suraksha",
    "photo",
    "voltaic",
    "margin",
    "tax",
    "deducted",
    "road",
    "lane",
    "showroom",
    "hospital",
    "house",
    "village",
    "taluka",
    "district",
    "industrial",
    "park",
    "complex",
    "baner",
    "bandra",
    "pune",
    "mumbai",
    "maharashtra",
    "deccan",
    "gymkhana",
    "shivaji",
    "nagar",
    "colony",
    "model",
    "chambers",
    "bapat",
    "marg",
    "erandawane",
    "khed",
    "chakan",
    "supa",
    "buena",
    "monte",
    "gopal",
    "bo",
    "opp",
    "opposite",
    "particulars",
    "date",
    "email",
    "telephone",
    "key",
    "managerial",
    "related",
    "widely",
    "circulated",
    "marathi",
    "daily",
    "newspaper",
    "non-gaap",
    "measures",
    "operational",
    "alpha",
    "bill",
    "delay",
    "fraud",
    "challan",
    "circular",
    "exchange",
    "excludes",
    "all",
    "offer-related",
    "individual",
    "qualified",
    "institutional",
    "buyer",
    "buyers",
    "qib",
    "qibs",
    "retail",
    "portion",
    "pre-offer",
    "post-offer",
    "selling",
    "shareholder",
    "secondary",
    "primary",
    "registered",
    "broker",
    "brokers",
    "upi",
    "mandate",
}

KNOWN_PERSON_NAMES = {
    "kushal hegde",
    "kushal subbayya hegde",
    "narayana b. shetty",
    "pushpa kushal hegde",
    "rajesh hegde",
    "rohit kushal hegde",
}

# High-confidence person names that spaCy may fail to detect.
EXPLICIT_PERSON_NAMES = {
    "rakhi girija shetty",
    "lokesh shah",
    "soumavo sarkar",
    "dinesh hirachand munot",
    "lalit muljibhai sarvaiya",
    "ganesh prasad",
}

# Specific geographic/address expressions that spaCy
# occasionally labels as PERSON.
PERSON_FALSE_POSITIVE_PHRASES = {
    "bandra",
    "bandra east",
    "bandra kurla",
    "bandra kurla complex",
    "baner",
    "baner pune",
    "bapat marg",
    "bapat marg pune",
    "bapat marg, pune",
    "buena monte",
    "chakan taluka",
    "chakan taluka - khed",
    "chakan taluka-khed",
    "deccan gymkhana",
    "deccan gymkhana society",
    "gopal bo",
    "gopal house",
    "gram jyoti",
    "kubera chambers opp",
    "listing sebi bhavan",
    "model colony",
    "road lane",
    "sancheti hospital shivajinagar",
    "shivaji nagar",
    "supa facility",
    "taluka khed",
    "taluka parner",
    "taluka - khed pune",
    "taluka-khed",
    "tanishq showroom",
    "tara chambers",
    "village khalumbre",
    "waterloo industrial",
    "waterloo industrial park",
    "waterloo industrial park ii private limited",
    "waterloo industrial park iii private limited",
    "waterloo industrial park iv private limited",
    "waterloo industrial park v private limited",
    "waterloo industrial park vi private limited",
    "waterloo industrial park viii private limited",
    "waterloo industrial park ix private limited",
    "waterloo industrial park ix b private limited",
}


# Words which strongly suggest that a multi-word entity
# is NOT a person's name.
NON_PERSON_PHRASE_WORDS = {
    "limited",
    "private",
    "ltd",
    "llp",
    "corporation",
    "bank",
    "securities",
    "industries",
    "fund",
    "trust",
    "family",
    "facility",
    "industrial",
    "park",
    "road",
    "street",
    "lane",
    "marg",
    "complex",
    "society",
    "hospital",
    "showroom",
    "branch",
    "house",
    "village",
    "taluka",
    "district",
    "pune",
    "mumbai",
    "maharashtra",
    "company",
    "department",
    "division",
    "committee",
    "board",
    "investor",
    "investors",
    "bidder",
    "bidders",
    "shareholder",
    "shareholders",
    "portion",
    "price",
    "amount",
    "offer",
    "bids",
    "bid",
    "funds",
    "mutual",
    "personnel",
    "managerial",
    "schedule",
    "reference",
    "rate",
}

def find_explicit_person_names(text):
    findings = []

    for name in EXPLICIT_PERSON_NAMES:

        pattern = re.compile(
            rf"(?<![A-Za-z])"
            rf"{re.escape(name)}"
            rf"(?![A-Za-z])",
            re.IGNORECASE,
        )

        for match in pattern.finditer(text):

            findings.append({
                "start": match.start(),
                "end": match.end(),
                "type": "person",
                "value": match.group(),
            })

    return findings


def clean_person_entity(entity_text):

    text = entity_text.strip()

    text = text.strip(
        " /,:;.-*^&()[]{}"
    )

    words = text.split()

    while words and (
        words[-1]
        .lower()
        .strip(" /,:;.-*^&()[]{}")
        in GENERIC_PERSON_WORDS
    ):

        words.pop()

    return " ".join(words)


def looks_like_person_name(name):

    if not name:

        return False

    name = clean_person_entity(
        name
    )

    if not name:

        return False

    lower_name = (
        re.sub(
            r"[^a-z\s]",
            " ",
            name.lower(),
        )
    )

    lower_name = re.sub(
        r"\s+",
        " ",
        lower_name,
    ).strip()

    if lower_name in PERSON_FALSE_POSITIVE_PHRASES:

        return False

    words = name.split()

    # A normal full name needs at least two words.
    if len(words) < 2:

        return False

    cleaned_words = []

    for word in words:

        cleaned = re.sub(
            r"[^A-Za-z'-]",
            "",
            word,
        )

        if cleaned:

            cleaned_words.append(
                cleaned
            )

    if len(cleaned_words) < 2:

        return False

    lower_words = {
        word.lower()
        for word in cleaned_words
    }

    # Reject obvious non-person phrases.
    if (
        lower_words
        & NON_PERSON_PHRASE_WORDS
    ):

        return False

    if (
        lower_words
        & GENERIC_PERSON_WORDS
    ):

        return False

    # Every remaining token must look like a name.
    for word in cleaned_words:

        if not re.fullmatch(
            r"[A-Za-z][A-Za-z'-]*",
            word,
        ):

            return False

    return True


# ============================================================
# ORGANIZATION FILTER
# ============================================================

ORGANIZATION_TERMS = (
    "limited",
    "ltd",
    "private limited",
    "pvt ltd",
    "corporation",
    "bank",
    "securities",
    "industries",
    "company",
    "llp",
    "trust",
)


GENERIC_ORGANIZATION_WORDS = {
    "company",
    "board",
    "offer",
    "registrar",
    "bank",
    "email",
    "website",
    "directors",
    "promoters",
    "bankers",
    "investors",
}


# ============================================================
# REPLACEMENT GENERATOR
# ============================================================

class ReplacementGenerator:

    def __init__(self):

        self.counters = {
            "email": 0,
            "phone": 0,
            "ip": 0,
            "credit_card": 0,
            "ssn": 0,
            "dob": 0,
            "person": 0,
            "organization": 0,
            "address": 0,
        }

        self.cache = {}

    def generate(
        self,
        pii_type,
        original,
    ):

        key = (
            pii_type,
            original.lower().strip(),
        )

        if key in self.cache:

            return self.cache[key]

        self.counters[
            pii_type
        ] += 1

        number = self.counters[
            pii_type
        ]

        replacements = {

            "email":
                f"person{number}@example.com",

            "phone":
                f"+91 90000 {number:05d}",

            "ip":
                f"192.0.2.{number}",

            "credit_card":
                f"4111-1111-1111-{number:04d}",

            "ssn":
                f"000-00-{number:04d}",

            "dob":
                "01 January 1990",

            "person":
                f"Person {number}",

            "organization":
                f"Example Organization {number}",

            "address":
                (
                    f"Example Address {number}, "
                    f"Pune, Maharashtra 400001"
                ),
        }

        replacement = replacements[
            pii_type
        ]

        self.cache[key] = replacement

        return replacement


# ============================================================
# VALIDATION
# ============================================================

def luhn_valid(number):

    digits = [
        int(c)
        for c in number
        if c.isdigit()
    ]

    if not 13 <= len(digits) <= 19:

        return False

    checksum = 0

    parity = len(digits) % 2

    for i, digit in enumerate(
        digits
    ):

        if i % 2 == parity:

            digit *= 2

            if digit > 9:

                digit -= 9

        checksum += digit

    return checksum % 10 == 0


def valid_phone(candidate):

    digits = re.sub(
        r"\D",
        "",
        candidate,
    )

    if not 10 <= len(digits) <= 13:

        return False

    try:

        number = phonenumbers.parse(
            candidate,
            "IN",
        )

        return (
            phonenumbers.is_possible_number(
                number
            )
            and
            phonenumbers.is_valid_number(
                number
            )
        )

    except phonenumbers.NumberParseException:

        return False


# ============================================================
# ALREADY REDACTED VALUES
# ============================================================

def is_generated_value(value):

    patterns = [

        r"person\d+@example\.com",

        r"\+91 90000 \d{5}",

        r"192\.0\.2\.\d+",

        r"4111-1111-1111-\d{4}",

        r"000-00-\d{4}",

        r"Person \d+",

        r"Example Organization \d+",

        r"Example Address \d+",
    ]

    return any(
        re.fullmatch(
            pattern,
            value.strip(),
            re.IGNORECASE,
        )
        for pattern in patterns
    )


# ============================================================
# DETECTION
# ============================================================

def detect_entities(
    text,
    nlp,
):

    findings = []

    def add(
        start,
        end,
        pii_type,
        value,
    ):

        if start >= end:

            return

        if is_generated_value(
            value
        ):

            return

        findings.append({
            "start": start,
            "end": end,
            "type": pii_type,
            "value": value,
        })

    # --------------------------------------------------------
    # EXPLICIT HIGH-CONFIDENCE PERSON NAMES
    # --------------------------------------------------------

    for finding in find_explicit_person_names(text):

        add(
            finding["start"],
            finding["end"],
            finding["type"],
            finding["value"],
        )

    # --------------------------------------------------------
    # EMAIL
    # --------------------------------------------------------

    for match in EMAIL_PATTERN.finditer(
        text
    ):

        add(
            match.start(),
            match.end(),
            "email",
            match.group(),
        )

    # --------------------------------------------------------
    # IP
    # --------------------------------------------------------

    for match in IP_PATTERN.finditer(
        text
    ):

        add(
            match.start(),
            match.end(),
            "ip",
            match.group(),
        )

    # --------------------------------------------------------
    # CREDIT CARD
    # --------------------------------------------------------

    for match in CREDIT_CARD_PATTERN.finditer(
        text
    ):

        if luhn_valid(
            match.group()
        ):

            add(
                match.start(),
                match.end(),
                "credit_card",
                match.group(),
            )

    # --------------------------------------------------------
    # SSN
    # --------------------------------------------------------

    for match in SSN_PATTERN.finditer(
        text
    ):

        add(
            match.start(),
            match.end(),
            "ssn",
            match.group(),
        )

    # --------------------------------------------------------
    # DOB
    # --------------------------------------------------------

    for match in DOB_PATTERN.finditer(
        text
    ):

        add(
            match.start(),
            match.end(),
            "dob",
            match.group(),
        )

    # --------------------------------------------------------
    # PHONE
    # --------------------------------------------------------

    for match in PHONE_PATTERN.finditer(
        text
    ):

        candidate = match.group().strip()

        if valid_phone(
            candidate
        ):

            add(
                match.start(),
                match.end(),
                "phone",
                candidate,
            )

    # --------------------------------------------------------
    # ADDRESS
    # --------------------------------------------------------

    for match in ADDRESS_LABEL_PATTERN.finditer(
        text
    ):

        start = match.start()

        search_end = min(
            len(text),
            match.end() + 500,
        )

        candidate = text[
            start:search_end
        ]

        pin = PIN_PATTERN.search(
            candidate
        )

        component = ADDRESS_COMPONENT_PATTERN.search(
            candidate
        )

        if pin and component:

            end = (
                start
                + pin.end()
            )

            add(
                start,
                end,
                "address",
                text[start:end],
            )

    # --------------------------------------------------------
    # KNOWN PERSON NAMES
    # --------------------------------------------------------
    # Directly detect known person names before spaCy.
    # This also handles names followed by footnote markers
    # such as * ^ &.

    for person_name in KNOWN_PERSON_NAMES:

        if not person_name:
            continue

        pattern = re.compile(
            r"(?<![A-Za-z])"
            + re.escape(person_name)
            + r"(?=$|[^A-Za-z])",
            re.IGNORECASE,
        )

        for match in pattern.finditer(text):

            add(
                match.start(),
                match.end(),
                "person",
                match.group(),
            )

    # --------------------------------------------------------
    # SPACY
    # --------------------------------------------------------

    # Process large paragraphs in smaller chunks.
    # This prevents spaCy from consuming too much memory
    # on low-memory deployment environments such as Render.
    SPACY_CHUNK_SIZE = 2000

    if len(text) <= SPACY_CHUNK_SIZE:
        spacy_chunks = [(0, nlp(text))]
    else:
        spacy_chunks = []

        for start in range(
            0,
            len(text),
            SPACY_CHUNK_SIZE,
        ):
            chunk = text[
                start:
                start + SPACY_CHUNK_SIZE
            ]

            if chunk.strip():
                spacy_chunks.append(
                    (
                        start,
                        nlp(chunk),
                    )
                )

    for offset, doc in spacy_chunks:

        for entity in doc.ents:

            # Convert spaCy's chunk-relative positions
            # back to positions in the original text.
            entity_start = (
                offset
                + entity.start_char
            )

            entity_end = (
                offset
                + entity.end_char
            )

            # ====================================================
            # PERSON
            # ====================================================

            if entity.label_ == "PERSON":

                value = clean_person_entity(
                    entity.text
                )

                normalized_value = (
                    value.strip().lower()
                )

                if (
                    normalized_value
                    not in KNOWN_PERSON_NAMES
                    and not looks_like_person_name(value)
                ):
                    continue

                # Check whether the entity is actually inside
                # an address-like region.
                nearby_left = text[
                    max(
                        0,
                        entity_start - 80,
                    ):
                    entity_start
                ].lower()

                nearby_right = text[
                    entity_end:
                    min(
                        len(text),
                        entity_end + 80,
                    )
                ].lower()

                nearby = (
                    nearby_left
                    + " "
                    + nearby_right
                )

                address_context_words = {
                    "road",
                    "marg",
                    "street",
                    "lane",
                    "floor",
                    "building",
                    "complex",
                    "hospital",
                    "showroom",
                    "village",
                    "taluka",
                    "district",
                    "nagar",
                    "society",
                    "branch",
                    "facility",
                    "pune",
                    "mumbai",
                    "maharashtra",
                }

                nearby_words = set(
                    re.findall(
                        r"[a-z]+",
                        nearby,
                    )
                )

                address_score = len(
                    nearby_words
                    & address_context_words
                )

                # If spaCy says PERSON but the surrounding
                # text strongly resembles an address, do not
                # classify it as a person.
                if address_score >= 3:
                    continue

                add(
                    entity_start,
                    entity_end,
                    "person",
                    value,
                )

            # ====================================================
            # ORGANIZATION
            # ====================================================

            elif entity.label_ == "ORG":

                value = entity.text.strip()

                lower = value.lower()

                if any(
                    term in lower
                    for term in ORGANIZATION_TERMS
                ):

                    if (
                        lower
                        not in GENERIC_ORGANIZATION_WORDS
                    ):

                        add(
                            entity_start,
                            entity_end,
                            "organization",
                            value,
                        )
    return findings


# ============================================================
# MERGE FINDINGS
# ============================================================

# ============================================================
# MERGE FINDINGS
# ============================================================

def merge_findings(
    findings
):

    findings = sorted(
        findings,
        key=lambda x: (
            x["start"],
            -(x["end"] - x["start"]),
        ),
    )

    result = []

    for finding in findings:

        if not result:

            result.append(
                finding
            )

            continue

        previous = result[-1]

        if (
            finding["start"]
            < previous["end"]
        ):

            previous_length = (
                previous["end"]
                - previous["start"]
            )

            current_length = (
                finding["end"]
                - finding["start"]
            )

            # If same start, prefer the more specific
            # entity type over a shorter generic entity.
            if (
                current_length
                > previous_length
            ):

                result[-1] = finding

        else:

            result.append(
                finding
            )

    return result


# ============================================================
# REDACT TEXT
# ============================================================

def redact_text(
    text,
    nlp,
    generator,
):

    findings = detect_entities(
        text,
        nlp,
    )

    findings = merge_findings(
        findings
    )

    if not findings:

        return text, []

    output = []

    cursor = 0

    for finding in findings:

        start = finding[
            "start"
        ]

        end = finding[
            "end"
        ]

        output.append(
            text[cursor:start]
        )

        replacement = generator.generate(
            finding["type"],
            finding["value"],
        )

        output.append(
            replacement
        )

        finding[
            "replacement"
        ] = replacement

        cursor = end

    output.append(
        text[cursor:]
    )

    return (
        "".join(output),
        findings,
    )

# ============================================================
# FINAL NAME SAFETY PASS
# ============================================================

FINAL_PERSON_NAMES = [
    "Kushal Hegde",
    "Kushal Subbayya Hegde",
    "Narayana B. Shetty",
    "Pushpa Kushal Hegde",
    "Rajesh Hegde",
    "Rohit Kushal Hegde",
]


def final_person_safety_pass(document, generator):

    replacements = []

    # DOCX package parts that can contain visible text
    parts = [
        document.part,
    ]

    # Main document XML
    root = document.part.element

    # Search every XML text node in the document body.
    for element in root.iter():

        if element.tag.endswith("}t"):

            if not element.text:
                continue

            original_text = element.text
            new_text = original_text

            for name in FINAL_PERSON_NAMES:

                pattern = re.compile(
                    re.escape(name),
                    re.IGNORECASE,
                )

                if pattern.search(new_text):

                    replacement = generator.generate(
                        "person",
                        name,
                    )

                    new_text = pattern.sub(
                        replacement,
                        new_text,
                    )

                    replacements.append({
                        "type": "person",
                        "original": name,
                        "replacement": replacement,
                        "location": "final XML safety pass",
                    })

            element.text = new_text

    return replacements

# ============================================================
# DOCUMENT PROCESSING
# ============================================================

def process_document():

    print(
        "Loading spaCy model..."
    )

    nlp = spacy.load(
        MODEL_NAME,
        disable=[
            "parser",
            "tagger",
            "lemmatizer",
            "attribute_ruler",
        ],
    )

    generator = (
        ReplacementGenerator()
    )

    document = Document(
        INPUT_FILE
    )

    all_findings = []

    # --------------------------------------------------------
    # PARAGRAPHS
    # --------------------------------------------------------

    for paragraph in document.paragraphs:

        text = paragraph.text

        if not text.strip():

            continue

        redacted, findings = redact_text(
            text,
            nlp,
            generator,
        )

        if findings:

            paragraph.text = redacted

            for finding in findings:

                all_findings.append({
                    **finding,
                    "location":
                        "paragraph",
                })

    # --------------------------------------------------------
    # TABLES
    # --------------------------------------------------------

    for table_number, table in enumerate(
        document.tables,
        start=1,
    ):

        for row_number, row in enumerate(
            table.rows,
            start=1,
        ):

            for column_number, cell in enumerate(
                row.cells,
                start=1,
            ):

                text = cell.text

                if not text.strip():

                    continue

                redacted, findings = redact_text(
                    text,
                    nlp,
                    generator,
                )

                if findings:

                    cell.text = redacted

                    for finding in findings:

                        all_findings.append({
                            **finding,
                            "location": (
                                f"table "
                                f"{table_number}, "
                                f"row "
                                f"{row_number}, "
                                f"column "
                                f"{column_number}"
                            ),
                        })

    # --------------------------------------------------------
    # FINAL PERSON SAFETY PASS
    # --------------------------------------------------------

    final_findings = final_person_safety_pass(
        document,
        generator,
    )

    all_findings.extend(
        final_findings
    )

    # --------------------------------------------------------
    # SAVE
    # --------------------------------------------------------

    OUTPUT_FILE.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document.save(
        OUTPUT_FILE
    )

    # --------------------------------------------------------
    # REPORT
    # --------------------------------------------------------

    counts = {}

    for finding in all_findings:

        pii_type = finding[
            "type"
        ]

        counts[pii_type] = (
            counts.get(
                pii_type,
                0,
            )
            + 1
        )

    print()
    print(
        "=" * 60
    )
    print(
        "PII REDACTION COMPLETE"
    )
    print(
        "=" * 60
    )

    print(
        f"Output: {OUTPUT_FILE}"
    )

    print(
        f"Total redactions: "
        f"{len(all_findings)}"
    )

    print()
    print("COUNTS")

    for pii_type, count in sorted(
        counts.items()
    ):

        print(
            f"{pii_type}: {count}"
        )

    print()
    print("REPLACEMENTS")

    for (
        (pii_type, original),
        replacement,
    ) in generator.cache.items():

        print(
            f"{pii_type}: "
            f"{original} -> "
            f"{replacement}"
        )


# ============================================================
# MAIN
# ============================================================

if __name__ == "__main__":

    process_document()