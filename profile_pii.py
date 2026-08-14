from pathlib import Path
import re
from collections import Counter

from docx import Document
import phonenumbers
from phonenumbers import NumberParseException


INPUT_FILE = Path("input/KSH_International_RHP.docx")
OUTPUT_FILE = Path("pii_inventory.txt")


# ============================================================
# REGEX PATTERNS
# ============================================================

EMAIL_PATTERN = re.compile(
    r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"
)

IP_PATTERN = re.compile(
    r"\b(?:\d{1,3}\.){3}\d{1,3}\b"
)

SSN_PATTERN = re.compile(
    r"\b\d{3}-\d{2}-\d{4}\b"
)

CREDIT_CARD_PATTERN = re.compile(
    r"\b(?:\d[ -]*?){13,19}\b"
)

DOB_PATTERN = re.compile(
    r"""
    (?:
        date\s+of\s+birth |
        dob |
        born\s+on |
        birth\s+date
    )
    [:\s-]{0,20}
    (
        \d{1,2}[/-]\d{1,2}[/-]\d{2,4}
        |
        \d{1,2}\s+
        (?:January|February|March|April|May|June|
        July|August|September|October|November|December)
        \s+\d{4}
    )
    """,
    re.IGNORECASE | re.VERBOSE,
)


# ============================================================
# PHONE DETECTION
# ============================================================

PHONE_CONTEXT = re.compile(
    r"""
    (?:
        telephone |
        phone |
        mobile |
        contact\s*(?:no|number)? |
        tel |
        fax
    )
    """,
    re.IGNORECASE | re.VERBOSE,
)


def is_likely_phone(value: str, surrounding_text: str) -> bool:
    """
    Use Google's libphonenumber database/rules to determine
    whether a numeric string is plausibly a telephone number.

    The RHP contains Indian numbers, so we primarily try
    region='IN'.
    """

    digits = re.sub(r"\D", "", value)

    # Reject clearly short values.
    if len(digits) < 7:
        return False

    # Reject values consisting entirely of zeros.
    if set(digits) == {"0"}:
        return False

    # Numbers such as 000013004 are likely document numbers,
    # not telephone numbers.
    if digits.startswith("000"):
        return False

    try:
        parsed = phonenumbers.parse(value, "IN")

        if phonenumbers.is_possible_number(parsed):
            if phonenumbers.is_valid_number(parsed):
                return True

    except NumberParseException:
        pass

    # Some Indian office numbers in the RHP are formatted in
    # ways libphonenumber may not recognize perfectly.
    # Context can therefore rescue a plausible number.
    if PHONE_CONTEXT.search(surrounding_text):

        # Indian telephone numbers normally contain 10 digits
        # excluding the country code.
        if len(digits) in {10, 11, 12, 13}:
            return True

    return False


# ============================================================
# CREDIT CARD VALIDATION
# ============================================================

def looks_like_credit_card(value: str) -> bool:

    digits = re.sub(r"\D", "", value)

    if not 13 <= len(digits) <= 19:
        return False

    # Luhn algorithm
    total = 0

    for index, digit in enumerate(reversed(digits)):

        number = int(digit)

        if index % 2 == 1:
            number *= 2

            if number > 9:
                number -= 9

        total += number

    return total % 10 == 0


# ============================================================
# IP VALIDATION
# ============================================================

def valid_ip(value: str) -> bool:

    try:
        parts = value.split(".")

        return (
            len(parts) == 4
            and all(
                0 <= int(part) <= 255
                for part in parts
            )
        )

    except ValueError:
        return False


# ============================================================
# DOCUMENT EXTRACTION
# ============================================================

def extract_text(document):

    chunks = []

    # -------------------------
    # Paragraphs
    # -------------------------

    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:
            chunks.append(
                ("paragraph", text)
            )

    # -------------------------
    # Tables
    # -------------------------

    for table_index, table in enumerate(
        document.tables,
        start=1
    ):

        for row_index, row in enumerate(
            table.rows,
            start=1
        ):

            for col_index, cell in enumerate(
                row.cells,
                start=1
            ):

                text = cell.text.strip()

                if text:

                    location = (
                        f"table {table_index}, "
                        f"row {row_index}, "
                        f"column {col_index}"
                    )

                    chunks.append(
                        (location, text)
                    )

    # -------------------------
    # Headers / Footers
    # -------------------------

    for section_index, section in enumerate(
        document.sections,
        start=1
    ):

        for paragraph in section.header.paragraphs:

            text = paragraph.text.strip()

            if text:

                chunks.append(
                    (
                        f"section {section_index} header",
                        text
                    )
                )

        for paragraph in section.footer.paragraphs:

            text = paragraph.text.strip()

            if text:

                chunks.append(
                    (
                        f"section {section_index} footer",
                        text
                    )
                )

    return chunks


# ============================================================
# MAIN
# ============================================================

def main():

    if not INPUT_FILE.exists():

        raise FileNotFoundError(
            f"Input document not found: {INPUT_FILE}"
        )

    document = Document(INPUT_FILE)

    chunks = extract_text(document)

    counters = Counter()
    findings = []

    for location, text in chunks:

        # ====================================================
        # EMAIL
        # ====================================================

        emails = EMAIL_PATTERN.findall(text)

        for value in emails:

            counters["email"] += 1

            findings.append(
                ("EMAIL", location, value)
            )

        # ====================================================
        # PHONE
        # ====================================================

        # Look for formatted numeric sequences rather than
        # treating arbitrary short numbers as phone numbers.

        possible_phone_pattern = re.compile(
            r"""
            (?:
                \+?\d{1,3}[\s.-]?
            )?
            (?:
                \(\d{2,5}\)
                |
                \d{2,5}
            )
            [\s.-]
            \d{3,4}
            [\s.-]
            \d{3,4}
            |
            \+?\d{10,13}
            """,
            re.VERBOSE
        )

        for match in possible_phone_pattern.finditer(text):

            value = match.group(0)

            if is_likely_phone(value, text):

                counters["phone"] += 1

                findings.append(
                    ("PHONE", location, value)
                )

        # ====================================================
        # IP ADDRESS
        # ====================================================

        for value in IP_PATTERN.findall(text):

            if valid_ip(value):

                counters["ip_address"] += 1

                findings.append(
                    ("IP_ADDRESS", location, value)
                )

        # ====================================================
        # SSN
        # ====================================================

        for value in SSN_PATTERN.findall(text):

            counters["ssn"] += 1

            findings.append(
                ("SSN", location, value)
            )

        # ====================================================
        # CREDIT CARD
        # ====================================================

        for value in CREDIT_CARD_PATTERN.findall(text):

            if looks_like_credit_card(value):

                counters["credit_card"] += 1

                findings.append(
                    ("CREDIT_CARD", location, value)
                )

        # ====================================================
        # DATE OF BIRTH
        # ====================================================

        for value in DOB_PATTERN.findall(text):

            counters["date_of_birth"] += 1

            findings.append(
                ("DATE_OF_BIRTH", location, value)
            )

    # ========================================================
    # WRITE REPORT
    # ========================================================

    with OUTPUT_FILE.open(
        "w",
        encoding="utf-8"
    ) as file:

        file.write("PII INVENTORY\n")
        file.write("=" * 70 + "\n\n")

        file.write(
            f"Input file: {INPUT_FILE}\n"
        )

        file.write(
            f"Text chunks scanned: {len(chunks)}\n\n"
        )

        file.write("COUNTS\n")
        file.write("-" * 70 + "\n")

        for category, count in sorted(
            counters.items()
        ):

            file.write(
                f"{category}: {count}\n"
            )

        file.write("\n\nFINDINGS\n")
        file.write("=" * 70 + "\n\n")

        for category, location, value in findings:

            file.write(
                f"[{category}] "
                f"{location}: "
                f"{value}\n"
            )

    print("=" * 60)
    print("PII PROFILING COMPLETE")
    print("=" * 60)

    print(
        f"Text chunks scanned: {len(chunks)}"
    )

    print("\nDetected candidates:")

    for category, count in sorted(
        counters.items()
    ):

        print(
            f"  {category}: {count}"
        )

    print(
        f"\nDetailed report: {OUTPUT_FILE}"
    )


if __name__ == "__main__":
    main()