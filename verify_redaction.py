from __future__ import annotations

import re
from pathlib import Path

import phonenumbers
import spacy
from docx import Document


# ============================================================
# CONFIGURATION
# ============================================================

INPUT_FILE = Path(
    "output/KSH_International_RHP_REDACTED.docx"
)

MODEL_NAME = "en_core_web_sm"


# ============================================================
# REGEX PATTERNS
# ============================================================

EMAIL_PATTERN = re.compile(
    r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"
)

IP_PATTERN = re.compile(
    r"\b(?:"
    r"(?:25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)\."
    r"){3}"
    r"(?:25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)\b"
)

CREDIT_CARD_PATTERN = re.compile(
    r"\b(?:"
    r"\d{4}[- ]?"
    r"){3}"
    r"\d{4}\b"
)

SSN_PATTERN = re.compile(
    r"\b\d{3}-\d{2}-\d{4}\b"
)

PHONE_PATTERN = re.compile(
    r"(?<!\d)"
    r"(?:\+?\d[\d\s().-]{8,}\d)"
    r"(?!\d)"
)


# ============================================================
# GENERATED VALUE PATTERNS
# ============================================================

GENERATED_PATTERNS = {

    "email": re.compile(
        r"person\d+@example\.com",
        re.IGNORECASE,
    ),

    "phone": re.compile(
        r"\+91 90000 \d{5}",
        re.IGNORECASE,
    ),

    "ip": re.compile(
        r"192\.0\.2\.\d+",
        re.IGNORECASE,
    ),

    "credit_card": re.compile(
        r"4111-1111-1111-\d{4}",
        re.IGNORECASE,
    ),

    "ssn": re.compile(
        r"000-00-\d{4}",
        re.IGNORECASE,
    ),

    "person": re.compile(
        r"Person \d+",
        re.IGNORECASE,
    ),

    "organization": re.compile(
        r"Example Organization \d+",
        re.IGNORECASE,
    ),

    "address": re.compile(
        r"Example Address \d+",
        re.IGNORECASE,
    ),
}


# ============================================================
# PHONE VALIDATION
# ============================================================

def valid_phone(candidate):

    digits = re.sub(
        r"\D",
        "",
        candidate,
    )

    if len(digits) < 10:
        return False

    try:

        number = phonenumbers.parse(
            candidate,
            "IN",
        )

        return (
            phonenumbers.is_possible_number(number)
            and
            phonenumbers.is_valid_number(number)
        )

    except phonenumbers.NumberParseException:

        return False


# ============================================================
# DOCUMENT TEXT EXTRACTION
# ============================================================

def extract_document_text(document):

    pieces = []

    # Paragraphs
    for paragraph in document.paragraphs:

        if paragraph.text.strip():

            pieces.append(
                paragraph.text
            )

    # Tables
    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                if cell.text.strip():

                    pieces.append(
                        cell.text
                    )

    return "\n".join(pieces)


# ============================================================
# MAIN VERIFICATION
# ============================================================

def main():

    print("=" * 60)
    print("PII REDACTION VERIFICATION")
    print("=" * 60)

    if not INPUT_FILE.exists():

        print(
            f"ERROR: File not found: {INPUT_FILE}"
        )

        return

    print(
        f"Checking: {INPUT_FILE}"
    )

    document = Document(
        INPUT_FILE
    )

    text = extract_document_text(
        document
    )

    print(
        f"Characters checked: {len(text):,}"
    )

    print()

    # --------------------------------------------------------
    # GENERATED PLACEHOLDERS
    # --------------------------------------------------------

    print("GENERATED PLACEHOLDERS")

    for pii_type, pattern in GENERATED_PATTERNS.items():

        matches = pattern.findall(text)

        print(
            f"{pii_type}: {len(matches)}"
        )

    print()

    # --------------------------------------------------------
    # POSSIBLE REMAINING EMAILS
    # --------------------------------------------------------

    print("POSSIBLE REMAINING EMAILS")

    emails = []

    for match in EMAIL_PATTERN.finditer(text):

        value = match.group()

        if not GENERATED_PATTERNS["email"].fullmatch(
            value
        ):

            emails.append(value)

    if emails:

        for email in sorted(set(emails)):

            print(
                f"  FOUND: {email}"
            )

    else:

        print(
            "  NONE"
        )

    print()

    # --------------------------------------------------------
    # POSSIBLE REMAINING PHONES
    # --------------------------------------------------------

    print("POSSIBLE REMAINING PHONE NUMBERS")

    phones = []

    for match in PHONE_PATTERN.finditer(text):

        candidate = match.group().strip()

        if not valid_phone(candidate):

            continue

        if GENERATED_PATTERNS["phone"].fullmatch(
            candidate
        ):

            continue

        phones.append(candidate)

    if phones:

        for phone in sorted(set(phones)):

            print(
                f"  FOUND: {phone}"
            )

    else:

        print(
            "  NONE"
        )

    print()

    # --------------------------------------------------------
    # IP ADDRESSES
    # --------------------------------------------------------

    print("POSSIBLE REMAINING IP ADDRESSES")

    ips = []

    for match in IP_PATTERN.finditer(text):

        value = match.group()

        if not GENERATED_PATTERNS["ip"].fullmatch(
            value
        ):

            ips.append(value)

    if ips:

        for ip in sorted(set(ips)):

            print(
                f"  FOUND: {ip}"
            )

    else:

        print(
            "  NONE"
        )

    print()

    # --------------------------------------------------------
    # CREDIT CARDS
    # --------------------------------------------------------

    print("POSSIBLE CREDIT CARD NUMBERS")

    cards = []

    for match in CREDIT_CARD_PATTERN.finditer(text):

        value = match.group()

        if not GENERATED_PATTERNS[
            "credit_card"
        ].fullmatch(value):

            cards.append(value)

    if cards:

        for card in sorted(set(cards)):

            print(
                f"  FOUND: {card}"
            )

    else:

        print(
            "  NONE"
        )

    print()

    # --------------------------------------------------------
    # SSN
    # --------------------------------------------------------

    print("POSSIBLE SSNs")

    ssns = []

    for match in SSN_PATTERN.finditer(text):

        value = match.group()

        if not GENERATED_PATTERNS[
            "ssn"
        ].fullmatch(value):

            ssns.append(value)

    if ssns:

        for ssn in sorted(set(ssns)):

            print(
                f"  FOUND: {ssn}"
            )

    else:

        print(
            "  NONE"
        )

    print()

    # --------------------------------------------------------
    # SPACY PERSON CHECK
    # --------------------------------------------------------

    print("POSSIBLE REMAINING PERSON ENTITIES")

    print(
        "Loading spaCy model..."
    )

    nlp = spacy.load(
        MODEL_NAME
    )

    doc = nlp(text)

    persons = []

    for entity in doc.ents:

        if entity.label_ != "PERSON":

            continue

        value = entity.text.strip()

        if not value:

            continue

        if GENERATED_PATTERNS["person"].fullmatch(
            value
        ):

            continue

        persons.append(value)

    persons = sorted(
        set(persons),
        key=str.lower,
    )

    if persons:

        for person in persons:

            print(
                f"  POSSIBLE: {person}"
            )

    else:

        print(
            "  NONE"
        )

    print()

    # --------------------------------------------------------
    # SPACY ORGANIZATION CHECK
    # --------------------------------------------------------

    print("POSSIBLE REMAINING ORGANIZATIONS")

    organizations = []

    for entity in doc.ents:

        if entity.label_ != "ORG":

            continue

        value = entity.text.strip()

        if not value:

            continue

        if GENERATED_PATTERNS[
            "organization"
        ].fullmatch(value):

            continue

        organizations.append(value)

    organizations = sorted(
        set(organizations),
        key=str.lower,
    )

    if organizations:

        for organization in organizations:

            print(
                f"  POSSIBLE: {organization}"
            )

    else:

        print(
            "  NONE"
        )

    print()

    # --------------------------------------------------------
    # FINAL RESULT
    # --------------------------------------------------------

    print("=" * 60)
    print("VERIFICATION COMPLETE")
    print("=" * 60)

    print()
    print(
        "IMPORTANT:"
    )
    print(
        "spaCy PERSON/ORG results are candidates, "
        "not automatically confirmed PII."
    )

    print(
        "Review those candidates before changing "
        "the redaction detector."
    )


# ============================================================
# ENTRY POINT
# ============================================================

if __name__ == "__main__":

    main()