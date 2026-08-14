from pathlib import Path
import re

from docx import Document


INPUT_FILE = Path("input/KSH_International_RHP.docx")
OUTPUT_FILE = Path("address_inventory.txt")


ADDRESS_CONTEXT = re.compile(
    r"""
    registered\s+office |
    corporate\s+office |
    registered\s+address |
    corporate\s+address |
    mailing\s+address |
    residential\s+address |
    correspondence\s+address |
    office\s+address |
    address
    """,
    re.IGNORECASE | re.VERBOSE,
)


PIN_PATTERN = re.compile(
    r"\b\d{6}\b"
)


def extract_text(document):

    chunks = []

    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:
            chunks.append(("paragraph", text))

    for table_index, table in enumerate(
        document.tables,
        start=1
    ):

        for row_index, row in enumerate(
            table.rows,
            start=1
        ):

            for col_index, cell in enumerate(
                row.cells,
                start=1
            ):

                text = cell.text.strip()

                if text:

                    location = (
                        f"table {table_index}, "
                        f"row {row_index}, "
                        f"column {col_index}"
                    )

                    chunks.append(
                        (location, text)
                    )

    for section_index, section in enumerate(
        document.sections,
        start=1
    ):

        for paragraph in section.header.paragraphs:

            text = paragraph.text.strip()

            if text:

                chunks.append(
                    (
                        f"section {section_index} header",
                        text
                    )
                )

        for paragraph in section.footer.paragraphs:

            text = paragraph.text.strip()

            if text:

                chunks.append(
                    (
                        f"section {section_index} footer",
                        text
                    )
                )

    return chunks


def main():

    document = Document(INPUT_FILE)

    chunks = extract_text(document)

    findings = []

    for location, text in chunks:

        # ----------------------------------------------------
        # Explicit address context
        # ----------------------------------------------------

        if ADDRESS_CONTEXT.search(text):

            findings.append(
                (
                    "EXPLICIT_ADDRESS_CONTEXT",
                    location,
                    text
                )
            )

        # ----------------------------------------------------
        # Indian PIN code
        # ----------------------------------------------------

        for pin in PIN_PATTERN.findall(text):

            findings.append(
                (
                    "PIN_CODE",
                    location,
                    pin
                )
            )

    with OUTPUT_FILE.open(
        "w",
        encoding="utf-8"
    ) as file:

        file.write("ADDRESS INVENTORY\n")
        file.write("=" * 70 + "\n\n")

        for category, location, value in findings:

            file.write(
                f"[{category}] "
                f"{location}: "
                f"{value}\n"
            )

    print("=" * 60)
    print("ADDRESS PROFILING COMPLETE")
    print("=" * 60)

    print(
        f"Candidates found: {len(findings)}"
    )

    print(
        f"Report: {OUTPUT_FILE}"
    )


if __name__ == "__main__":
    main()