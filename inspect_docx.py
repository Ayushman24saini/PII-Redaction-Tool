from pathlib import Path

from docx import Document


INPUT_FILE = Path("input/KSH_International_RHP.docx")


def inspect_document(path: Path) -> None:
    if not path.exists():
        raise FileNotFoundError(f"Input file not found: {path}")

    document = Document(path)

    print("=" * 60)
    print("DOCX INSPECTION")
    print("=" * 60)

    print(f"File: {path}")
    print(f"Paragraphs: {len(document.paragraphs)}")
    print(f"Tables: {len(document.tables)}")
    print(f"Sections: {len(document.sections)}")

    print("\n--- FIRST 20 NON-EMPTY PARAGRAPHS ---")

    count = 0

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if not text:
            continue

        print(f"{count + 1}: {text[:300]}")

        count += 1

        if count >= 20:
            break

    print("\n--- TABLE INFORMATION ---")

    for index, table in enumerate(document.tables[:10], start=1):
        print(
            f"Table {index}: "
            f"{len(table.rows)} rows x {len(table.columns)} columns"
        )


if __name__ == "__main__":
    inspect_document(INPUT_FILE)